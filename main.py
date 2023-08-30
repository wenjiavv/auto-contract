import docx
import os
import config
import tempfile
import shutil
from docx2pdf import convert
from datetime import datetime

TEMPLATE_PATH = config.CONFIG['TEMPLATE_PATH']
OUTPUT_PATH = config.CONFIG['OUTPUT_PATH']

for filename in os.listdir(TEMPLATE_PATH):
    if not filename.endswith('.docx'):
        continue
    template_file = os.path.join(TEMPLATE_PATH, filename)

    doc = docx.Document(template_file)

    elements = config.elements

    for element, value in elements.items():
        for paragraph in doc.paragraphs:
            if element in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if element in inline[i].text:
                        inline[i].text = inline[i].text.replace(element, value)
                        print(f"替换了变量 {element}，新值为 {value}")

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if element in cell.text:
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            inline = paragraph.runs
                            for i in range(len(inline)):
                                if element in inline[i].text:
                                    inline[i].text = inline[i].text.replace(element, value)
                                    print(f"替换了变量 {element}，新值为 {value}")

    temp_dir = tempfile.mkdtemp()
    temp_file = os.path.join(temp_dir, filename)
    doc.save(temp_file)

    # 读取临时文件
    doc = docx.Document(temp_file)

    # 再次替换模版中的要素
    for element, value in elements.items():
        for paragraph in doc.paragraphs:
            if element in paragraph.text:
                paragraph.text = paragraph.text.replace(element, value)
                print(f"第二次替换了变量 {element}，新值为 {value}")

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if element in cell.text:
                        cell.text = cell.text.replace(element, value)
                        print(f"第二次替换了变量 {element}，新值为 {value}")




    contract_month = elements['[month]']
    contract_day = elements['[day]']
    new_filename = filename.replace('TEMP', f'{contract_month}{contract_day}')
    output_file = os.path.join(OUTPUT_PATH, new_filename)
    doc.save(output_file)
    shutil.rmtree(temp_dir)

    pdf_file = os.path.splitext(output_file)[0] + ".pdf"
    convert(output_file, pdf_file)

    os.remove(output_file)

today = datetime.today().strftime('%Y-%m-%d')

# 创建新的文件夹名字
new_dir_name = "保理合同" + today

# 指定新文件夹路径
dst_dir = os.path.join(os.path.dirname(OUTPUT_PATH), new_dir_name)

# 复制并重命名文件夹
shutil.copytree(OUTPUT_PATH, dst_dir)

# 创建 zip 文件
shutil.make_archive(dst_dir, 'zip', dst_dir)

print("操作完成。")
